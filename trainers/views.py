from django.http import HttpResponse
from django.template import loader
from django.shortcuts import render,redirect,get_object_or_404
from django.utils.timezone import datetime
from django.db.models import Sum, F, Count
from .models import *
import json
from django.contrib import messages
from django.utils import timezone
import calendar
from datetime import timedelta
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from dateutil.relativedelta import relativedelta
from .middleware import require_organization
from decimal import Decimal




from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
import json

@login_required(login_url='/login/')
@require_organization
def Home(request):        
    organization = request.organization
    staff = Staff.objects.get(user=request.user)
    if staff.is_admin:

        now = timezone.now()
        today = now.date()
        # Financial summaries
        paymentss = Payments.objects.filter(organization=organization)
        financial_summary = {
            'total_yearly_income': paymentss.filter(
                paymentdate__year=today.year
            ).aggregate(total=Sum('paymentAmount'))['total'] or 0,
            'monthly_income': paymentss.filter(
                paymentdate__year=today.year,
                paymentdate__month=today.month
            ).aggregate(total=Sum('paymentAmount'))['total'] or 0,
            'today_income': paymentss.filter(
                paymentdate=today
            ).aggregate(total=Sum('paymentAmount'))['total'] or 0
        }

        payment_categories = {
            'month': {
                'label': 'شهرية', 
                'frequency': 'monthly', 
                'grace_days': 0
            },
            'subscription': {
                'label': 'انخراط', 
                'frequency': 'yearly', 
                'grace_days': 0
            },
            'assurance': {
                'label': 'التأمين', 
                'frequency': 'yearly', 
                'grace_days': 0
            },'jawaz': {
                'label': 'جواز', 
                'frequency': 'yearly', 
                'grace_days': 0
            }
        }

        payment_status = {}

        for category, category_info in payment_categories.items():
            trainers = Trainer.objects.filter(is_active=True, organization=organization)
            unpaid_trainers = []

            for trainer in trainers:
                # Retrieve the last payment for the trainer in this category
                last_payment = paymentss.filter(
                    trainer=trainer,
                    paymentCategry=category
                ).order_by('-paymentdate').first()

                if last_payment:

                    payment_due_date = None
                    
                    if category_info['frequency'] == 'monthly':
                        payment_due_date = last_payment.paymentdate + relativedelta(months=1)
                        
                    
                    elif category_info['frequency'] == 'yearly':
                        payment_due_date = last_payment.paymentdate.replace(
                            year=last_payment.paymentdate.year + 1
                        ) + timedelta(days=category_info['grace_days'])

                    # Check if payment is overdue
                    if today >= payment_due_date:
                        unpaid_trainers.append({
                            'trainer_id': trainer.id,  # Added trainer ID
                            'trainer_name': f"{trainer.first_name} {trainer.last_name}",
                            'last_payment_date': last_payment.paymentdate
                        })
                        
                        
                else:
                    # Trainer has never paid in this category
                    unpaid_trainers.append({
                        'trainer_id': trainer.id,  # Added trainer ID
                        'trainer_name': f"{trainer.first_name} {trainer.last_name}",
                        'last_payment_date': None
                    })
            
            payment_status[category] = {
                'label': category_info['label'],
                'unpaid_trainers': unpaid_trainers,
                'total_unpaid_trainers': len(unpaid_trainers)
            }
        # Paid today trainees
        paid_today_trainees = paymentss.filter(paymentdate=today).select_related('trainer')
        paid_today_trainees = [
            {
                "trainer_name": f"{payment.trainer.first_name} {payment.trainer.last_name}",
                "payment_date": payment.paymentdate,
                "payment_category": payment.get_paymentCategry_display(),
                "payment_amount": payment.paymentAmount
            }
            for payment in paid_today_trainees
        ]

        # Chart data
        chart_labels = [str(m) for m in range(1, 13)]  # Months 1-12
        chart_data = {category: [0] * 12 for category in payment_categories}

        for category in payment_categories:
            category_income = paymentss.filter(
                paymentCategry=category,
                paymentdate__year=today.year
            ).values('paymentdate__month').annotate(
                total_income=Sum('paymentAmount')
            )

            for entry in category_income:
                month = entry['paymentdate__month'] - 1
                chart_data[category][month] = entry['total_income']

        safe_chart_data = {
        key: json.dumps(value, default=float)   # <--- FIXED HERE
        for key, value in chart_data.items()
    }
        # Prepare context
        context = {
            'financial_summary': financial_summary,  # Financial overview
            'chart_labels': json.dumps(chart_labels),  # Chart labels
            'chart_data': safe_chart_data,  # Chart data
            'payment_status': payment_status,  # Payment tracking details
            'paid_today_trainees': paid_today_trainees,  # Trainees who paid today
            
        }

        return render(request, "pages/index.html", context)
    else:
        return redirect('dashboard')


def landing_page(request):
    return render(request, "pages/landing_page.html")


@csrf_exempt
@require_POST
@require_organization
def bulk_deactivate_trainers(request):
    organization = request.organization
    staff = Staff.objects.get(user=request.user)
    """
    View to handle bulk deactivation of trainers
    """
    if not staff.is_admin:
        return JsonResponse({'success': False, 'error': 'Unauthorized'}, status=403)
    
    try:
        data = json.loads(request.body)
        trainer_ids = data.get('trainer_ids', [])
        
        if not trainer_ids:
            return JsonResponse({'success': False, 'error': 'No trainers selected'})
        
        # Update trainers to inactive
        updated_count = Trainer.objects.filter(
            id__in=trainer_ids,
            is_active=True,
             organization=organization
        ).update(is_active=False)
        
        return JsonResponse({
            'success': True, 
            'message': f'تم إلغاء تفعيل {updated_count} متدرب بنجاح',
            'deactivated_count': updated_count
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})



"""
def Home(request):
    # Get today's date
    today = datetime.now().date()
    today_day = today.day  # Day of the month

    # Determine the last day of the current month
    _, last_day_of_month = calendar.monthrange(today.year, today.month)

    # Lists for trainees
    paid_today_trainees = []
    due_today_trainees = []
    total_year_monthes = 0
    total_month =0
    newpayment=0
    newtrainer=0

    for trainee in Trainer.objects.all():
        # Get all payments for this trainee
        trainee_payments = Payments.objects.filter(trainer=trainee,paymentCategry="month").order_by("-paymentdate")
        
        if trainee_payments.exists():
            # Get the most recent payment
            if trainee_payments.first().paymentdate.month == today.month and trainee_payments.first().paymentdate.year==today.year:
                total_month += trainee_payments.first().paymentAmount
            last_payment = trainee_payments.first()
            for pay in trainee_payments:
                total_year_monthes += pay.paymentAmount

            if last_payment.paymentdate == today:
                # Trainee has paid today
                newpayment += last_payment.paymentAmount
                if trainee.started_day == today:newtrainer += 1
                
                paid_today_trainees.append({
                    "trainee_name": trainee.last_name,
                    "last_payment_date": last_payment.paymentdate,
                })
            else:
                payment_due_day = min(last_payment.paymentdate.day, last_day_of_month)
                if (
                    payment_due_day <= today_day and
                    (last_payment.paymentdate.month < today.month or last_payment.paymentdate.year < today.year)
                ):
                    due_today_trainees.append({
                        "trainee_name": trainee.last_name,
                        "last_payment_date": last_payment.paymentdate,
                    })
        else:
            # If no payments exist, trainee is due today
            due_today_trainees.append({
                "trainee_name": trainee.last_name,
                "last_payment_date": None,
            })



    monthly_data = (
        Payments.objects.annotate(month=TruncMonth('paymentdate'))
        .values('month')
        .annotate(total=Sum('paymentAmount'))
        .order_by('month')
    )

    # Prepare the data for the chart
    chart_labels = [entry['month'].strftime('%Y-%m') for entry in monthly_data]
    chart_data = [entry['total'] for entry in monthly_data]

    
            

    return render(request, "pages/index.html", {
        "paid_today_trainees": paid_today_trainees,
        "due_today_trainees": due_today_trainees,
        "totalyear" : total_year_monthes,
        "totalmonth" : total_month,
        "newpayment":newpayment,
        "newtrainer":newtrainer,
        'chart_labels': chart_labels,
        'chart_data': chart_data,
    })

"""

def addme(request, org_slug):
    """
    Public registration page - organization identified by URL slug
    URL: /register/nojoum-arkana/
    """
    organization = get_object_or_404(OrganizationInfo, slug=org_slug)
    
    # Check if organization has reached trainer limit
    if organization.trainers.filter(is_active=True).count() >= organization.max_trainers:
        return render(request, "pages/registration_closed.html", {
            'organization': organization,
            'message': 'عذراً، لقد وصلت الجمعية للحد الأقصى من المتدربين'
        })
    
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        birthday = request.POST.get('birthday')
        gender = request.POST.get('gender')
        phone = request.POST.get('phone', '0') or ''
        phone_parent = request.POST.get('phone_parent', '0') or ''
        email = request.POST.get('email', 'None@emale.com')
        address = request.POST.get('address', 'None')
        cin = request.POST.get('cin', 'None')
        education = request.POST.get('education')
        belt = request.POST.get('belt')
        upload = request.FILES.get('upload')
        category = request.POST.get('category')
        height = request.POST.get('height') or 0
        weight = request.POST.get('weight') or 0
        
        # Get document files
        national_id_doc = request.FILES.get('national_id_doc')
        civil_status_doc = request.FILES.get('civil_status_doc')

        if first_name and last_name and birthday and gender and education and category and cin:
            # Check if trainer already exists in THIS organization
            if Trainer.objects.filter(
                organization=organization,
                first_name=first_name,
                last_name=last_name
            ).exists():
                messages.error(request, "المتدرب موجود بالفعل في النظام.")
                return render(request, "pages/addme.html", {'organization': organization})
            
            # Create trainer for this specific organization
            trainer = Trainer.objects.create(
                organization=organization,
                first_name=first_name,
                last_name=last_name,
                birth_day=birthday,
                phone=phone,
                phone_parent=phone_parent,
                email=email,
                address=address,
                CIN=cin,
                male_female=gender,
                belt_degree=belt,
                Degree=education,
                category=category,
                started_day=datetime.today(),
                image=upload,
                tall=height,
                weight=weight
            )
            
            # Handle document uploads (optional)
            if national_id_doc:
                TrainerDocument.objects.create(
                    trainer=trainer,
                    document_type='بطاقة الوطنية',
                    file=national_id_doc
                )
            
            if civil_status_doc:
                TrainerDocument.objects.create(
                    trainer=trainer,
                    document_type='الحالة المدنية',
                    file=civil_status_doc
                )
            
            message = f"تمت إضافة المتدرب {first_name} إلى {organization.name} بنجاح"
            return render(request, "pages/sucss.html", {
                "org": organization
            })
        else:
            messages.error(request, "الرجاء ملء جميع الحقول المطلوبة.")
    
    return render(request, "pages/addme.html", {'organization': organization})

def addmedone(request):
    """Success page after registration"""
    return render(request, "pages/sucss.html")

@login_required(login_url='/login/')
@require_organization
def export_data(request, category):
    organization = request.organization
    if category == "all":
        trainees = Trainer.objects.filter(organization=organization)
    else:
        trainees = Trainer.objects.filter(organization=organization, category=category)

    # Create a CSV response
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="trainees_{category}.csv"'

    writer = csv.writer(response)
    writer.writerow(['First Name', 'Last Name', 'Email', 'Phone', 'Category'])
    for trainee in trainees:
        writer.writerow([trainee.first_name, trainee.last_name, trainee.email, trainee.phone, trainee.category])

    return response
@login_required(login_url='/login/')
def add_trainee(request):
    organization = request.organization
    
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        birthday = request.POST.get('birthday')
        gender = request.POST.get('gender')
        phone = request.POST.get('phone', 0) 
        phone_parent = request.POST.get('phone_parent', 0) or 0
        email = request.POST.get('email', 'None@emale.com')
        address = request.POST.get('address', 'Argana')
        cin = request.POST.get('cin', 'لا يوجد')
        education = request.POST.get('education')
        belt = request.POST.get('belt')
        upload = request.FILES.get('upload')
        category = request.POST.get('category')
        height = request.POST.get('height', 0) or 0
        weight = request.POST.get('weight', 0) or 0
        
        # Get document files
        national_id_doc = request.FILES.get('national_id_doc')
        civil_status_doc = request.FILES.get('civil_status_doc')

        if first_name and last_name and birthday and gender and education and category:
            if Trainer.objects.filter(organization=organization, first_name=first_name, last_name=last_name).exists():
                messages.error(request, "المتدرب موجود بالفعل.")
                return render(request, "pages/add_trainee.html")
            else:
                # Create trainer
                trainer = Trainer.objects.create(
                    organization=organization,
                    first_name=first_name,
                    last_name=last_name,
                    birth_day=birthday,
                    phone=phone,
                    email=email,
                    address=address,
                    CIN=cin,
                    male_female=gender,
                    belt_degree=belt,
                    Degree=education,
                    category=category,
                    started_day=datetime.today(),
                    image=upload,
                    tall=height,
                    weight=weight,
                    phone_parent=phone_parent
                )
                
                # Handle document uploads (optional)
                if national_id_doc:
                    TrainerDocument.objects.create(
                        trainer=trainer,
                        document_type='بطاقة الوطنية',
                        file=national_id_doc
                    )
                
                if civil_status_doc:
                    TrainerDocument.objects.create(
                        trainer=trainer,
                        document_type='الحالة المدنية',
                        file=civil_status_doc
                    )
                
                messages.success(request, "تمت إضافة المتدرب " + first_name + " بنجاح")
        else:
            messages.error(request, "فشلت عملية الإضافة.")
    
    if 'add_women' in request.path:
        return render(request, "pages/add_women.html")
    return render(request, "pages/add_trainee.html")

@login_required(login_url='/login/')
@require_organization
def add_payment(request):
        organization=request.organization
        if request.method == 'POST':
            trainer_id = request.POST.get('trainer')
            paymentdate = request.POST.get('paymentdate')
            paymentCategry = request.POST.get('paymentCategry')
            paymentAmount = request.POST.get('paymentAmount')

            # إنشاء كائن جديد للدفع
            trainer = Trainer.objects.get(id=trainer_id)
            payment = Payments(
                organization=organization,
                trainer=trainer,
                paymentdate=paymentdate,
                paymentCategry=paymentCategry,
                paymentAmount=paymentAmount
            )
            payment.save()
            
            return redirect('added_payment')

        trainers = Trainer.objects.filter(organization=organization).order_by('-id')
        return render(request, 'pages/add_payment.html', {'trainers': trainers})

@login_required(login_url='/login/')
@require_organization
def added_payment(request):
        
        return render(request,"pages/added_payment.html")

@login_required(login_url='/login/')
@require_organization
def payments_history(request):
        organization = request.organization
        from django.core.paginator import Paginator

        # Get search query from GET parameters
        search_query = request.GET.get('search', '').strip()
        
        # Start with all payments
        payments = Payments.objects.select_related('trainer').filter(organization=organization)
        
        # Apply search filter if query exists
        if search_query:
            from django.db.models.functions import Concat

            payments = payments.annotate(
                trainer_full_name=Concat(
                    'trainer__first_name',
                    Value(' '),
                    'trainer__last_name'
                )
            ).filter(
                Q(trainer__first_name__icontains=search_query) |
                Q(trainer__last_name__icontains=search_query) |
                Q(trainer_full_name__icontains=search_query) |
                Q(paymentCategry__icontains=search_query)
            )
        
        # Order by most recent
        payments = payments.order_by('-id')
        
        # Paginate: 25 items per page (adjust as needed)
        paginator = Paginator(payments, 25)
        page_number = request.GET.get('page', 1)
        page_obj = paginator.get_page(page_number)
        
        context = {
            'payments': page_obj,
            'search_query': search_query,
            'paginator': paginator,
            'page_obj': page_obj
        }
        return render(request, "pages/payments_history.html", context)
    
@login_required(login_url='/login/')
@require_organization
def payments_del(request,id):
        org=request.organization
        Payments.objects.get(organization=org, pk=id).delete()
        return redirect("payments_history")

@login_required(login_url='/login/')
@require_organization
def payment_edit(request, id):
    # استرجاع الدفعة باستخدام معرّفها
        org=request.organization
        payment = get_object_or_404(Payments, organization=org, id=id)

        if request.method == "POST":
            # تحديث البيانات الواردة من النموذج
            payment.paymentCategry = request.POST.get("paymentCategry")
            payment.paymentdate = request.POST.get("paymentdate")
            payment.paymentAmount = request.POST.get("paymentAmount")
            
            # حفظ التعديلات
            payment.save()
            #messages.success(request, "تم تحديث بيانات الدفعة بنجاح!")
            return redirect("payments_history")  # تعديل وجهة الإعادة كما هو مطلوب

        # إذا كانت الطلبية GET، عرض النموذج مع البيانات الحالية
        return render(request, "pages/edit_payment.html", {"payment": payment})

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from django.http import HttpResponse



def add_table_borders(table):
    tbl = table._element
    borders = parse_xml(r'''
    <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>
    ''')
    tbl.tblPr.append(borders)


def clone_table_style(src_table, dst_table):
    dst_table.style = src_table.style
    dst_table.autofit = src_table.autofit


def copy_header_row(src_table, dst_table):

    src_row = src_table.rows[0]
    dst_row = dst_table.rows[0]

    for i in range(len(src_row.cells)):
        dst_cell = dst_row.cells[i]
        dst_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Clear default content and use bold text
        paragraph = dst_cell.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run(src_row.cells[i].text)
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def update_word_table(word_file_path, payments, payment_category='assurance'):
    """
    Open Word document, clear existing table, and populate with new data.
    Starts a new table every 18 rows, each on a new page, preserving original table style and headers.
    """
    doc = Document(word_file_path)

    # Get the original table and its style
    original_table = doc.tables[0]

    # Clear all rows except the header
    for i in range(len(original_table.rows) - 1, 0, -1):
        original_table._element.remove(original_table.rows[i]._element)

    table = original_table
    add_table_borders(table)

    for i, p in enumerate(payments, start=1):
        if i > 1 and (i) % 16 == 0:
            # start a new page
            # Add new table
            table = doc.add_table(rows=1, cols=5) 
            table.autofit = original_table.autofit
            clone_table_style(original_table, table)
            add_table_borders(table)
            copy_header_row(original_table, table)

        row = table.add_row()
        
        cell_texts = [
            str(i),
            f"{p.trainer.last_name}",
            f"{p.trainer.first_name}",
            f" ",
            f"{p.trainer.birth_day}",
        ]

        for cell, text in zip(row.cells, cell_texts):
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Vertical centering
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal centering

    # Prepare the Word response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="payments_{payment_category}_{datetime.today()}.docx"'
    doc.save(response)

    return response


import xlwt
from django.http import HttpResponse

@login_required(login_url='/login/')
@require_organization
def export_xls(request):
    organization = request.organization
    payment_category = request.GET.get("category")
    start_date = request.GET.get("start_date")
    end_date = request.GET.get("end_date")
    trainer_category = request.GET.get("trainer_category")

    payments = Payments.objects.select_related('trainer').all()


    
    if payment_category and payment_category != "all":
        payments = payments.filter(organization=organization, paymentCategry=payment_category)
    
    if start_date and start_date != "None":
        payments = payments.filter(organization=organization, paymentdate__gte=start_date)
        if end_date and end_date != "None":
            payments = payments.filter(organization=organization, paymentdate__lte=end_date)
    if payment_category == 'assurance':
        current_dir = os.path.dirname(os.path.abspath(__file__))
        word_file_path = os.path.join(current_dir, "NOJOUM ARGANA ASSURANCE  N°14.docx")
        
        # Update the Word document
        return update_word_table(word_file_path, payments, payment_category)
        
        
    else:
        # Define header style
        header_style = xlwt.XFStyle()
        header_font = xlwt.Font()
        header_font.bold = True
        header_style.font = header_font
        wb = xlwt.Workbook(encoding='utf-8')
        ws = wb.add_sheet('Payments')
    
        # Define header style
        header_style = xlwt.XFStyle()
        header_font = xlwt.Font()
        header_font.bold = True
        header_style.font = header_font
       
        # Headers
        columns = ['المدرب', 'فئة المتدرب', 'تاريخ الدفع', 'نوع الدفع', 'المبلغ']
        for col_num, column_title in enumerate(columns):
            ws.write(0, col_num, column_title, header_style)

        # Data rows
        for row_num, p in enumerate(payments, start=1):
            ws.write(row_num, 0, f"{p.trainer.first_name} {p.trainer.last_name}")
            ws.write(row_num, 1, p.trainer.get_category_display())
            ws.write(row_num, 2, p.paymentdate.strftime("%Y-%m-%d"))
            ws.write(row_num, 3, p.get_paymentCategry_display())
            ws.write(row_num, 4, p.paymentAmount)
        
        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = f'attachment; filename="payments_{datetime.today()}.xls"'

        wb.save(response)
        return response
    
@login_required(login_url='/login/')
@require_organization
def export_csv(request):
    organization = request.organization
    payment_category = request.GET.get("category")
    start_date = request.GET.get("start_date")
    end_date = request.GET.get("end_date")
    trainer_category = request.GET.get("trainer_category")

    payments = Payments.objects.select_related('trainer').all()

    if payment_category:
        payments = payments.filter(organization=organization, paymentCategry=payment_category)
    if start_date:
        payments = payments.filter(organization=organization, paymentdate__gte=start_date)
    if end_date:
        payments = payments.filter(organization=organization, paymentdate__lte=end_date)
    if trainer_category:
        payments = payments.filter(organization=organization, trainer__category=trainer_category)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = "attachment; filename=payments.csv"

    writer = csv.writer(response)
    writer.writerow(["المدرب", "فئة المتدرب", "تاريخ الدفع", "نوع الدفع", "المبلغ"])

    for p in payments:
        writer.writerow([
            f"{p.trainer.first_name} {p.trainer.last_name}",
            p.trainer.get_category_display(),
            p.paymentdate,
            p.get_paymentCategry_display(),
            p.paymentAmount
        ])

    return response


import csv

@login_required
@require_organization
def export_csv(request):
    organization = request.organization
    payment_category = request.GET.get("category")
    start_date = request.GET.get("start_date")
    end_date = request.GET.get("end_date")
    trainer_category = request.GET.get("trainer_category")

    payments = Payments.objects.select_related('trainer').filter(organization=organization)

    if payment_category:
        payments = payments.filter(organization=organization,paymentCategry=payment_category)
    if start_date:
        payments = payments.filter(organization=organization,paymentdate__gte=start_date)
    if end_date:
        payments = payments.filter(organization=organization,paymentdate__lte=end_date)
    if trainer_category:
        payments = payments.filter(organization=organization,trainer__category=trainer_category)

    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = "attachment; filename=payments.csv"

    writer = csv.writer(response)
    writer.writerow(["المدرب", "فئة المتدرب", "تاريخ الدفع", "نوع الدفع", "المبلغ"])

    for p in payments:
        writer.writerow([
            f"{p.trainer.first_name} {p.trainer.last_name}",
            p.trainer.get_category_display(),
            p.paymentdate,
            p.get_paymentCategry_display(),
            p.paymentAmount
        ])

    return response

@login_required(login_url='/login/')
@require_organization
def download_documents(request):
    organization = request.organization
    # Get filter values from query params
    payment_category = request.GET.get("category")
    start_date = request.GET.get("start_date")
    end_date = request.GET.get("end_date")
    trainer_category = request.GET.get("trainer_category")

    payments = Payments.objects.select_related('trainer').filter(organization=organization)

    if payment_category:
        payments = payments.filter(organization=organization,paymentCategry=payment_category)

    if start_date:
        payments = payments.filter(organization=organization,paymentdate__gte=start_date)

    if end_date:
        payments = payments.filter(organization=organization,paymentdate__lte=end_date)

    if trainer_category:
        payments = payments.filter(organization=organization,trainer__category=trainer_category)

    context = {
        "payments": payments,
        "categories": Payments.get_catchoices(),
        "trainer_categories": Trainer.CatChoices,
        "selected_category": payment_category,
        "selected_trainer_category": trainer_category,
        "start_date": start_date,
        "end_date": end_date,
    }
    return render(request, "pages/download_documents.html", context)




from django.db.models import Q,Value
@login_required(login_url='/login/')
@require_organization
def trainees(request,category):
    organization = request.organization
    trainers = Trainer.objects.filter(organization=organization)
    if category!="all":
        trainers = trainers.filter(category=category)
    else:
        trainers = trainers.exclude(category="women")
    if request.method == "GET":
        if 'gender' in request.GET and request.GET['gender']:trainers= trainers.filter(male_female= request.GET['gender'])
        if 'order' in  request.GET:
            if request.GET['order']=='first_first':trainers= trainers.order_by('-started_day')
            if request.GET['order']=='last_first':trainers= trainers.order_by('started_day')
            if request.GET['order']=='first_name':trainers= trainers.order_by('last_name')
                    
    
    template = loader.get_template('pages/olders.html')
    return HttpResponse(template.render({'trainers':trainers,'number':trainers.count()},request))


@login_required(login_url='/login/')
@require_organization
def add_pay_from_prof(request,id,category,date):
    organization = request.organization
    if category == 'jawaz':
        amount = 150
        Payments.objects.create(organization=organization,trainer = Trainer.objects.get(pk=id),paymentCategry=category,paymentAmount=amount,paymentdate=date).save()
    elif category == 'assurance':
        amount = 100
        Payments.objects.create(organization=organization,trainer = Trainer.objects.get(pk=id),paymentCategry=category,paymentAmount=amount,paymentdate=date).save()
    elif category == 'subscription':
        amount = 50
        Payments.objects.create(organization=organization,trainer = Trainer.objects.get(pk=id),paymentCategry=category,paymentAmount=amount,paymentdate=date).save()
    elif category == 'month':
        amount = 100
        Payments.objects.create(organization=organization,trainer = Trainer.objects.get(pk=id),paymentCategry=category,paymentAmount=amount,paymentdate=date).save()
    return redirect("profile",id)

@login_required(login_url='/login/')
@require_organization
def trainee_profile(request, id):
    organization = request.organization
    trainer = get_object_or_404(Trainer, pk=id, organization=organization)
    
    if request.method == 'POST':
        # Handle document upload
        if request.POST.get('action') == 'upload_document':
            document_type = request.POST.get('document_type')
            document_file = request.FILES.get('document_file')
            
            if document_type and document_file:
                # Check if document type already exists
                existing_doc = TrainerDocument.objects.filter(
                    trainer=trainer,
                    document_type=document_type
                ).first()
                
                if existing_doc:
                    # Update existing document
                    existing_doc.file = document_file
                    existing_doc.save()
                    messages.success(request, f"تم تحديث {document_type} بنجاح")
                else:
                    # Create new document
                    TrainerDocument.objects.create(
                        trainer=trainer,
                        document_type=document_type,
                        file=document_file
                    )
                    messages.success(request, f"تم رفع {document_type} بنجاح")
            else:
                messages.error(request, "الرجاء اختيار نوع الوثيقة والملف")
        
        # Handle document deletion
        elif request.POST.get('action') == 'delete_document':
            doc_id = request.POST.get('doc_id')
            try:
                doc = TrainerDocument.objects.get(id=doc_id, trainer=trainer)
                doc_type = doc.document_type
                doc.file.delete()  # Delete the file from storage
                doc.delete()  # Delete the database record
                messages.success(request, f"تم حذف {doc_type} بنجاح")
            except TrainerDocument.DoesNotExist:
                messages.error(request, "الوثيقة غير موجودة")
        
        # Handle payment addition
        elif 'paymentCategry[]' in request.POST and request.POST.get('paymentdate'):
            payment_categories = request.POST.getlist('paymentCategry[]')
            payment_date = request.POST['paymentdate']
            
            for category in payment_categories:
                add_pay_from_prof(request, id, category, payment_date)
            
            messages.success(request, "تمت إضافة الدفعات بنجاح")
    
    # Get all payments for the trainee
    trainee_payments = Payments.objects.filter(organization=organization, trainer_id=id).order_by("paymentdate")
    articles = Article.objects.filter(organization=organization, trainees=trainer)
    docs = TrainerDocument.objects.filter(trainer_id=id).order_by('-uploaded_at')
    
    # Determine the range of months
    if trainee_payments.exists():
        first_payment_date = trainee_payments.first().paymentdate
        start_month = datetime(first_payment_date.year, first_payment_date.month, 1)
    else:
        start_month = datetime(datetime.now().year, datetime.now().month, 1)

    # Generate months from the first payment to the current month
    current_month = datetime.now()
    months = []
    while start_month <= current_month:
        months.append(start_month.strftime("%Y-%m"))
        if start_month.month == 12:
            start_month = datetime(start_month.year + 1, 1, 1)
        else:
            start_month = datetime(start_month.year, start_month.month + 1, 1)
    
    jawaz = Payments.objects.filter(trainer_id=id, paymentCategry='jawaz')
    assurance = Payments.objects.filter(trainer_id=id, paymentCategry='assurance')
    subscription = Payments.objects.filter(trainer_id=id, paymentCategry='subscription')
    
    # Track paid months
    paid_months = {payment.paymentdate.strftime("%Y-%m") for payment in trainee_payments.filter(paymentCategry='month')}

    # Prepare context with payment status
    payment_status = [
        {"month": month, "status": "paid" if month in paid_months else "unpaid"}
        for month in months
    ]

    return render(request, "pages/profile.html", {
        "payment_status": payment_status,
        "trainers": trainer,
        'articles': articles,
        'jawaz': jawaz,
        'assurance': assurance,
        'subscription': subscription,
        'docs': docs
    })
@login_required(login_url='/login/')
@require_organization
def delete_trainer_view(request, id):
    organization = request.organization
    # استرجاع المدرب وحذفه
    trainer = get_object_or_404(Trainer, id=id, organization=organization)
    cat = trainer.category
    trainer.delete()
    messages.success(request, "تم حذف المدرب بنجاح.")
    return redirect("trainees",'all')  # استبدل "trainers_list" بمسار القائمة الرئيسية


@login_required(login_url='/login/')
@require_organization
def edit_trainee(request, id):
    organization = request.organization
    trainee = get_object_or_404(Trainer, id=id, organization=organization)
    if request.method == 'POST':
        trainee.first_name = request.POST['first_name']
        trainee.last_name = request.POST['last_name']
        trainee.birth_day = request.POST['birthday']
        trainee.male_female = request.POST['gender']
        trainee.email = request.POST['email']
        trainee.phone = request.POST['phone']
        trainee.phone_parent = request.POST['phone_parent']
        trainee.address = request.POST['address']
        trainee.belt_degree = request.POST['belt']
        trainee.category = request.POST['category']
        trainee.weight = request.POST['weight']
        trainee.tall = request.POST['height']
        trainee.CIN = request.POST['cin']
        trainee.is_active = 'is_active' in request.POST
        if 'upload' in request.FILES:
            trainee.image = request.FILES['upload']
        trainee.save()
        return redirect("profile",trainee.pk)  # توجيه إلى قائمة المتدربين
    return render(request, 'pages/edit_trainee.html', {'trainee': trainee})




def success(request):
    return render(request, 'pages/sucss.html')




#JAWAZ
@login_required(login_url='/login/')
@require_organization
def add_article(request):
    organization = request.organization
    if request.method == "POST":
        title = request.POST['title']
        date = request.POST.get('date', timezone.now())
        content = request.POST['content']
        area = request.POST.get('area')
        category = request.POST.get('category')
        trainees_ids = request.POST.getlist('trainees')
        participetion_price = request.POST['payed']
        costs = request.POST['costs']
        location = request.POST.get('location')

        # إنشاء المقال
        article = Article.objects.create(
            organization=organization,
            title=title,
            area=area,
            category=category,
            date=date,
            content=content,
            costs=costs,
            location=location,
            participetion_price=participetion_price
        )

        # ربط المدربين بالمقال
        if trainees_ids:
            trainees = Trainer.objects.filter(organization=organization,id__in=trainees_ids)
            article.trainees.set(trainees)

        return redirect('articles/all')  # وجه إلى صفحة المقالات بعد الحفظ

    # بيانات المدربين لإظهارها في القائمة
    trainers = Trainer.objects.filter(organization=organization)
    return render(request, 'pages/add_Article.htm', {'trainers': trainers, 'date': timezone.now().date()})



@login_required(login_url='/login/')
@require_organization
def articles(request,category):
    organization = request.organization
    if category != 'all':
        articles = Article.objects.filter(category=category, organization=organization)
    else:
        articles = Article.objects.filter(organization=organization)
    context = {
        'articles':articles,
        'number':articles.count()
    }
    return render(request,'pages/articles.html',context)

def article_details(request,id):
    article = Article.objects.get(pk=id)
    
    con = {'article' : article
           }
    return render(request,'pages/article_detail.html',con)


@login_required(login_url='/login/')
@require_organization
def remove_article(request,id):
    organization = request.organization
    Article.objects.get(pk=id, organization=organization).delete()
    return redirect('articles','all')

#Financil
from calendar import monthrange
from datetime import datetime, timedelta

@login_required(login_url='/login/')
@require_organization
def finantial_status(request):
    organization = request.organization
    dm = timezone.datetime
    # Retrieve start and end dates from GET parameters
    start = request.GET.get('start', '2025-01-01')
    end = request.GET.get('end', '2025-12-31')
    
    # Convert string dates to datetime objects and then to date objects
    start_date = datetime.strptime(start, "%Y-%m-%d").date()
    end_date = datetime.strptime(end, "%Y-%m-%d").date()
    
    # Initialize totals
    rent_total = 0
    staff_total = 0
    rent_count = 0
    # Calculate recurring rent payments
    organization_info = organization  # Get organization info
    today = datetime.today().date()
    if organization_info.datepay:
        current_date = organization_info.datepay

        # Count the number of times the rent day occurs in the range
        
        
        # Format the date as YYYY MM DD.
        today_date = datetime.strptime(str(today), "%Y-%m-%d").date()
        while current_date <= end_date and current_date <= today_date:
            if current_date >= start_date:
                rent_count += 1
            # Move to the next month
            days_in_month = monthrange(current_date.year, current_date.month)[1]
            current_date += timedelta(days=days_in_month)

        rent_total = rent_count * organization_info.rent_amount

    # Calculate recurring staff payments
    staff = Staff.objects.filter(organization=organization)
    
    staff_totals = []
    for staff_member in staff:
        pay_day = staff_member.datepay.day
        current_date = staff_member.started
        # Format the date as YYYY MM DD.
        today_date =datetime.strptime(str(today),"%Y-%m-%d").date() 
        # Count the number of times the pay day occurs in the range
        pay_count = 0
        while current_date <= end_date and current_date <= today_date:
            if current_date >= start_date:
                
                pay_count += 1
            # Move to the next month
            if current_date.month == 12:
                current_date = current_date.replace(year=current_date.year + 1, month=1)
            else:
                days_in_month = monthrange(current_date.year, current_date.month)[1]
                current_date += timedelta(days=days_in_month)
        # Calculate the total salary for this staff member
        total_salary = pay_count * staff_member.salary
        staff_total += total_salary
        if total_salary==0:continue
        staff_totals.append({
        "staff_member":staff_member,
        "total_salary":total_salary,
        "pay_count":pay_count
        })
    # Calculate other costs and profits
    date__range = [start, end]
    payments = Payments.objects.filter(organization=organization,paymentdate__range=date__range)
    costs = Costs.objects.filter(organization=organization,date__range=date__range)
    articles = Article.objects.filter(organization=organization,date__range=date__range)
    total_added_costs = sum([x.amount for x in costs])
    total_article_costs = sum([x.costs for x in articles])
    total_costs = total_added_costs + total_article_costs + rent_total + staff_total
    arts_pro = sum([x.profit for x in articles])
    added_payments = Addedpay.objects.filter(organization=organization,date__range=date__range)
    total_added_pay = sum([x.amount for x in added_payments])

    payments_total = sum([  x.paymentAmount  for x in payments  ])
    profit = Decimal(payments_total + total_added_pay+ arts_pro)
    
    # Count payers by category
    payers_nbs = {
        'big': payments.filter(trainer_id__category='big').count(),
        'med': payments.filter(trainer_id__category='med').count(),
        'small': payments.filter(trainer_id__category='small').count(),
        'women': payments.filter(trainer_id__category='women').count(),
    }

    # Prepare context
    context = {
        'costs': costs,
        'total_costs': total_costs,
        'profit': profit,
        'net_profit': profit - total_costs,
        'numbers': payers_nbs,
        'articles_nb': articles.count(),
        'artspro': arts_pro,
        'payments_total': payments_total,
        'addedpayments': added_payments,
        'total_added_pay': total_added_pay,
        'added_expenses': costs,
        'total_added_costs': total_added_costs,
        'total_article_costs': total_article_costs,
        'staff': staff_totals,
        'staff_total': staff_total,  # Total salaries
        'pay_count': pay_count,
        'rent_count': rent_count,
        'rent_total': rent_total,  # Total rent
        'org': organization_info,
        'start': start,
        'end': end,
    }


    return render(request,'pages/finantial_status.html',context)

"""
import openai

def generate_ai_insights(data):
    openai.api_key = "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"
    
    prompt = f"Provide observations and insights based on the following financial data: {data}"
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        #=[
            {"role": "system", "content": "You are a financial advisor."},
            {"role": "user", "content": prompt}
        ]
    )
    return response['choices'][0]['message']['content']

    # Print the assistant's reply
    """
    

@login_required(login_url='/login/')
@require_organization
def add_payments(request):
    organization = request.organization
    if request.method == "POST":
        # Handle form submission directly from POST data
        
            title = request.POST.get("title")
            date = request.POST.get("date")
            description = request.POST.get("description")
            amount = float(request.POST.get("amount"))

            # Create a new expense entry
            Addedpay(
                organization=organization,
                date=date,
                desc=description,
                amount=amount,
                title=title
            ).save()
            return redirect("finantial_status")  # Redirect back to the same page after saving
    else:
            return render(
                request,
                "pages/add_payments.html",
       )


@login_required(login_url='/login/')
@require_organization
def add_expenses(request):
    organization = request.organization
    if request.method == "POST":
        # Handle form submission directly from POST data
        
            title = request.POST.get("title")
            date = request.POST.get("date")
            description = request.POST.get("description")
            amount = float(request.POST.get("amount"))

            # Create a new expense entry
            Costs(
                organization=organization,
                date=date,
                desc=description,
                amount=amount,
                cost=title
            ).save()
            return redirect("finantial_status")  # Redirect back to the same page after saving
    else:
            return render(
                request,
                "pages/add_expenses.html",
       )
    #return render(request, "pages/add_expenses.html")

@login_required(login_url='/login/')
@require_organization
def expenses_history(request):
    organization = request.organization
    return render(request, "pages/added_expences.html", {'payments': Costs.objects.filter(organization=organization)})

@login_required(login_url='/login/')
@require_organization
def delete_expense(request, id):
    organization = request.organization
    Costs.objects.filter(pk=id, organization=organization).delete()
    return redirect("expenses_history")

@login_required(login_url='/login/')
@require_organization
def edit_expense(request, id):
    organization = request.organization
    expense = get_object_or_404(Costs, id=id, organization=organization)

    if request.method == 'POST':
        expense.date = request.POST.get('date')
        expense.desc = request.POST.get('description')
        expense.amount = request.POST.get('amount')
        expense.cost = request.POST.get('title')
        expense.save()
        messages.success(request, "تم تحديث بيانات المصاريف بنجاح")
        return redirect('expenses_history')

    return render(request, 'pages/edit_expense.html', {'expense': expense})

@login_required(login_url='/login/')
@require_organization
def added_payments_history(request):
    organization = request.organization
    return render(request, "pages/added_payments_history.html", {'payments': Addedpay.objects.filter(organization=organization)})

@login_required(login_url='/login/')
@require_organization
def delete_pay(request, id):
    organization = request.organization
    Addedpay.objects.filter(pk=id, organization=organization).delete()
    return redirect("added_payments_history")


## MAnagin staff

@login_required(login_url='/login/')
@require_organization
def add_staff(request):
    try:
        staff = Staff.objects.get(user=request.user)
        if not staff.is_admin:
            messages.error(request, 'You do not have permission to add staff members.')
            return redirect('staff_list')

        organization = request.organization

        if request.method == "POST":
            try:
                # Get form data
                username = request.POST.get('username')
                email = request.POST.get('email')
                password = request.POST.get('password')
                role = request.POST.get('role')
                salary = float(request.POST.get('salary', 0))
                is_admin = request.POST.get('is_admin') == 'true'
                date = request.POST.get('date', timezone.now())

                # Validate required fields
                if not all([username, email, password]):
                    messages.error(request, 'All required fields must be filled.')
                    return render(request, 'pages/add_staff.html')

                # Check for existing username/email
                if User.objects.filter(username=username).exists():
                    messages.error(request, 'Username already exists.')
                    return render(request, 'pages/add_staff.html')

                if User.objects.filter(email=email).exists():
                    messages.error(request, 'Email already exists.')
                    return render(request, 'pages/add_staff.html')

                # Create user and staff in transaction
                with transaction.atomic():
                    user = User.objects.create_user(
                        username=username,
                        email=email,
                        password=password
                    )

                    Staff.objects.create(
                        organization=organization,
                        user=user,
                        role=role,
                        is_admin=is_admin,
                        salary=salary,
                        started=date
                    )

                messages.success(request, 'تم بنجاح إضافة موظف جديد.')
                return redirect('staff_list')

            except Exception as e:
                messages.error(request, f'Error adding staff member: {str(e)}')
                return render(request, 'pages/add_staff.html')

        return render(request, 'pages/add_staff.html')

    except Staff.DoesNotExist:
        messages.error(request, 'Staff record not found.')
        return redirect('login')

#infos

@login_required(login_url='/login/')
@require_organization
def staff_list(request):
    organization = request.organization
     # التحقق من صلاحيات المستخدم
    

    if Staff.objects.filter(organization=organization,user=request.user).exists():pass
    else : Staff.objects.create(
                organization=organization,
                user=request.user,
                role='ADMIN',
                is_admin=True,
                salary=0,
                started = timezone.now()
            )

     # جلب قائمة الموظفين
    staff_members = Staff.objects.filter(organization=organization)

    return render(request, 'pages/staff_list.html', {
        'organization': organization,
        'staff_members': staff_members,
    })


@login_required(login_url='/login/')
@require_organization
def edit_staff(request, staff_id):
    organization = request.organization
    staff = get_object_or_404(Staff, id=staff_id, organization=organization)

    if request.method == 'POST':
        staff.user.username = request.POST.get('username')
        staff.role = request.POST.get('role')
        staff.is_admin = request.POST.get('is_admin') == 'true'
        staff.started = request.POST.get('started')
        staff.salary = float(request.POST.get('salary'))

        staff.user.save()
        staff.save()

        #.success(request, "تم تحديث بيانات الموظف بنجاح")
        return redirect('staff_list')

    return render(request, 'pages/edit_staff.html', {'staff': staff})


@login_required(login_url='/login/')
@require_organization
def delete_staff(request, staff_id):
    organization = request.organization
    staff = get_object_or_404(Staff, id=staff_id, organization=organization)
    if staff.is_admin:
            return redirect('home')
        
    staff.user.delete()  # يحذف المستخدم المرتبط
    staff.delete()  # يحذف سجل الموظف

    #.success(request, "تم حذف الموظف بنجاح")
    return redirect('staff_list')


@login_required(login_url='/login/')
@require_organization
def edit_organization(request):
        organization = request.organization
        if request.method == 'POST':
            if organization:
                organization.name = request.POST.get('name')
                organization.description = request.POST.get('description')
                organization.established_date = request.POST.get('established_date')
                organization.rent_amount = float(request.POST.get('rent_amount'))
                organization.phone_number = request.POST.get('phone_number')
                organization.email = request.POST.get('email')
                organization.datepay = request.POST.get('payrentdate')
                organization.save()

            else:OrganizationInfo.objects.create(name = request.POST.get('name')
                ,description = request.POST.get('description')
                ,established_date = request.POST.get('established_date')
                ,rent_amount = float(request.POST.get('rent_amount'))
                ,phone_number = request.POST.get('phone_number')
                ,email = request.POST.get('email')
                ,datepay = request.POST.get('payrentdate')
                ).save()
            

        
            #.success(request, "تم تحديث بيانات الجمعية بنجاح")
            return redirect('staff_list')

        return render(request, 'pages/edit_organization.html', {'organization': organization})
        


## LOGIN


def login_view(request):
    # If already logged in, redirect appropriately
    if request.user.is_authenticated:
        try:
            staff = Staff.objects.select_related('organization').get(user=request.user)
            # User has organization
            if staff.is_admin:
                return redirect('home')
            else:
                return redirect('dashboard')
        except Staff.DoesNotExist:
            # User logged in but no organization
            if request.user.is_superuser:
                return redirect('home')  # Superuser can access home
            else:
                return redirect('setup_organization')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        # Authenticate the user
        user = authenticate(request, username=username, password=password)
        
        if user:
            # Login first
            login(request, user)
            
            # Now check if the user has a staff record
            try:
                staff = Staff.objects.select_related('organization').get(user=user)
                
                # Check if organization is active
                if not staff.organization.is_active:
                    messages.error(request, 'حساب الجمعية غير نشط. الرجاء التواصل مع الإدارة.')
                    logout(request)
                    return render(request, 'pages/login.html')
                
                messages.success(request, 'تم تسجيل الدخول بنجاح')
                
                # Redirect based on user type
                
                return redirect('home')
                
                    
            except Staff.DoesNotExist:
                # User exists but no staff record
                if user.is_superuser:
                    # Superuser without organization - allow access to home
                    messages.info(request, 'مرحباً! يمكنك إنشاء جمعية جديدة')
                    return redirect('home')
                else:
                    # Regular user without organization - redirect to setup
                    messages.info(request, 'يرجى إنشاء جمعية جديدة')
                    return redirect('setup_organization')
        else:
            messages.error(request, 'اسم المستخدم أو كلمة المرور غير صحيحة')
            return render(request, 'pages/login.html')
    
    return render(request, 'pages/login.html')

@login_required(login_url='/login/')
def logout_view(request):
    if request.user.is_authenticated:
        logout(request)
        messages.success(request, 'تم تسجيل الخروج بنجاح')
        return redirect('login')


@login_required(login_url='/login/')
@require_organization
def dashboard(request):
    organization = request.organization
    # Basic staff dashboard view
    try:
        """staff_profile = request.user.staff
        context = {
            'staff': staff_profile
        }"""
        now = timezone.now()
        today = now.date()

        payment_categories = {
            'month': {
                'label': 'شهرية', 
                'frequency': 'monthly', 
                'grace_days': 0
            },
            'subscription': {
                'label': 'انخراط', 
                'frequency': 'yearly', 
                'grace_days': 0
            },
            'assurance': {
                'label': 'التأمين', 
                'frequency': 'yearly', 
                'grace_days': 0
            }
        }

        
        payment_status = {}

        for category, category_info in payment_categories.items():
            trainers = Trainer.objects.filter(organization=organization, is_active=True)
            unpaid_trainers = []

            for trainer in trainers:
                # Retrieve the last payment for the trainer in this category
                last_payment = Payments.objects.filter(
                    organization=organization,
                    trainer=trainer,
                    paymentCategry=category
                ).order_by('-paymentdate').first()

                if last_payment:
                    payment_due_date = None
                    
                    if category_info['frequency'] == 'monthly':
                        # Get the last day of the month for the last payment
                        last_day_of_payment_month = calendar.monthrange(
                            last_payment.paymentdate.year, 
                            last_payment.paymentdate.month
                        )[1]  # E.g., 28, 29, 30, or 31
                        payment_due_date = last_payment.paymentdate.replace(
                            day=min(last_payment.paymentdate.day,last_day_of_payment_month),
                            month=today.month,
                            year=today.year
                        )
                    
                    elif category_info['frequency'] == 'yearly':
                        payment_due_date = last_payment.paymentdate.replace(
                            year=last_payment.paymentdate.year + 1
                        ) + timedelta(days=category_info['grace_days'])

                    # Check if payment is overdue
                    if today > payment_due_date:
                        unpaid_trainers.append({
                            'trainer_name': f"{trainer.first_name} {trainer.last_name}",
                            'last_payment_date': last_payment.paymentdate
                        })
                        
                        
                else:
                    # Trainer has never paid in this category
                    unpaid_trainers.append({
                        'trainer_name': f"{trainer.first_name} {trainer.last_name}",
                        'last_payment_date': None
                    })
            
            payment_status[category] = {
                'label': category_info['label'],
                'unpaid_trainers': unpaid_trainers,
                'total_unpaid_trainers': len(unpaid_trainers)
            }
        # Paid today trainees
        paid_today_trainees = Payments.objects.filter(paymentdate=today).select_related('trainer')
        paid_today_trainees = [
            {
                "trainer_name": f"{payment.trainer.first_name} {payment.trainer.last_name}",
                "payment_date": payment.paymentdate,
                "payment_category": payment.get_paymentCategry_display(),
                "payment_amount": payment.paymentAmount
            }
            for payment in paid_today_trainees
        ]

        for category in payment_categories:
            category_income = Payments.objects.filter(
                organization=organization,
                paymentCategry=category,
                paymentdate__year=today.year
            ).values('paymentdate__month').annotate(
                total_income=Sum('paymentAmount')
            )

            for entry in category_income:
                month = entry['paymentdate__month'] - 1

        # Prepare context
            context = {
                'payment_status': payment_status,  # Payment tracking details
                'paid_today_trainees': paid_today_trainees,  # Trainees who paid today
            }

            return render(request, 'pages/home.html',context)
    except:
        messages.error(request, 'الوصول غير مسموح')
        return redirect('login')
#Email
from django.core.mail import send_mail

def emails(request):
    con={
        'emails':Emailed.objects.all()
    }

    return render(request,'pages/emails.html',con)

"""
from background_task import background

@background
def send_payment_reminder():
    today = timezone.now().date()
    print("=================================================================\n")
    print(today)
    # Payment categories with labels and frequencies
    payment_categories = {
        'month': {'label': 'الشهر', 'frequency': 'monthly'},
        'subscription': {'label': 'الانخراط', 'frequency': 'yearly'},
        'assurance': {'label': 'التأمين', 'frequency': 'yearly_A'},
    }

    for category, category_info in payment_categories.items():
        trainers = Trainer.objects.filter(is_active=True)

        for trainer in trainers:
            # Fetch the most recent payment
            last_payment = Payments.objects.filter(
                trainer=trainer,
                paymentCategry=category
            ).order_by('-paymentdate').first()

            if last_payment:
                start_date = last_payment.paymentdate  # Last payment date

                # Determine the next payment due date
                if category_info['frequency'] == 'monthly':
                    due_day = start_date.day

                    # Handle end-of-month payments
                    if due_day == calendar.monthrange(start_date.year, start_date.month)[1]:
                        due_day = calendar.monthrange(today.year, today.month)[1]

                    due_date = today.replace(day=due_day) if today.day <= due_day else None
                else:  # Yearly payments (subscription, assurance)
                    due_date = start_date.replace(year=today.year)

                # Ensure the due_date is valid and today matches it
                if not due_date or today != due_date:
                    continue
 
                # Check if an email has already been sent for this user, category, and day
                if not Emailed.objects.filter(
                    user=trainer,
                    category=category_info['label'],
                    datetime=today
                ).exists():
                    # Prepare the email body
                    email_subject = f"تذكير بدفع مستحقات {category_info['label']}"
                    email_message = (
                        f"عزيزي {trainer.first_name} {trainer.last_name},\n\n" if trainer.male_female=="male" else f"عزيزتي {trainer.first_name} {trainer.last_name},\n\n" 
                        
                        f"نود تذكيرك بأن موعد دفع مستحقات  '{category_info['label']}' "
                        f"يوافق اليوم ({due_date}).\n"
                        f"الرجاء إتمام الدفع.\n\n"
                        f"إذا كنت بحاجة إلى أي مساعدة، لا تتردد في التواصل مع الادارة.\n\n"
                        f"شكرًا لك،\n"
                        f"إدارة نجوم أركانة"
                    )
 
                    # Send the reminder email
                    send_mail(
                        subject=email_subject,
                        message=email_message,
                        from_email="youness.bouhnif.84@edu.uiz.ac.ma",
                        recipient_list=[trainer.email],
                        fail_silently=False,
                    )

                    # Log the email in the database
                    Emailed.objects.create(
                        user=trainer,
                        email=trainer.email,
                        category=category_info['label'],
                        datetime=timezone.now(),
                    )
                    print(f"Reminder email sent to {trainer.email} for category {category_info['label']}")
                else:
                    print(f"Email already sent to {trainer.email} for category {category_info['label']} today.")
                print("======================================================")

from trainers.tasks import send_payment_reminder_task

# Schedule the task to run daily
send_payment_reminder_task(repeat=60*60*24)

#handling errors



from django.http import JsonResponse
"""
@login_required 
@require_organization
def edit_article(request, id):
    org = request.organization
    article = get_object_or_404(Article,organization=org, id=id)
    trainees = Trainer.objects.filter(organization=org)

    if request.method == 'POST':
        
            # Update basic article information
            article.title = request.POST.get('title')
            article.date = request.POST.get('date')
            article.location = request.POST.get('location')
            article.participetion_price = float(request.POST.get('profitpayed', 0))
            article.costs = float(request.POST.get('costs', 0))
            article.content = request.POST.get('content')
            article.organization = org

                        
            # Update trainees
            selected_trainees = request.POST.getlist('trainees')
            article.trainees.clear()
            article.trainees.add(*selected_trainees)
            
            # Save the article
            article.save()
            
            messages.success(request, 'تم تحديث المقالة بنجاح')
            return redirect('article_detail', id=article.id)
            
       
    
    context = {
        'article': article,
        'trainees_art': article.trainees.all(),
        'trainees': trainees,
    }
    
    return render(request, 'pages/edit_article.html', context)


import pandas as pd


@login_required(login_url='/login/')
@require_organization
def export_data(request,category):
    organization = request.organization
    data = None
    if category == 'payments':
        data = Payments.objects.filter(organization=organization).values()
    elif category == 'trainers':
        data = Trainer.objects.filter(organization=organization).values()
    elif category == 'articles':
        data = Article.objects.filter(organization=organization).values()
    elif category == 'staff':
        data = Staff.objects.filter(organization=organization).values()
    elif category == 'added_payments':
        data = Addedpay.objects.filter(organization=organization).values()
    elif category == 'expenses':
        data = Costs.objects.filter(organization=organization).values()
    else:
        messages.error(request, 'الفئة غير معروفة')
    # Query the Person model to get all records
    if data:
        # Convert the QuerySet to a DataFrame
        df = pd.DataFrame(list(data))
        # Define the Excel file response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename={category}.xlsx'
        # Use Pandas to write the DataFrame to an Excel file
        df.to_excel(response, index=False, engine='openpyxl')
        return response


from datetime import datetime, timedelta
@login_required(login_url='/login/')
@require_organization
def unpaid_trainees(request):
    # Generate a list of years (e.g., from 2020 to the current year)
    current_year = datetime.now().year
    years = range(2020, current_year + 1)
    org = request.organization
    # List of months
    months = [
        (1, 'يناير'), (2, 'فبراير'), (3, 'مارس'), (4, 'أبريل'),
(5, 'مايو'), (6, 'يونيو'), (7, 'يوليو'), (8, 'أغسطس'),
(9, 'سبتمبر'), (10, 'أكتوبر'), (11, 'نوفمبر'), (12, 'ديسمبر')
    ]

    if request.method == 'GET':
        year = request.GET.get('year')
        month = request.GET.get('month')
        cat = request.GET.get('category','all')

        if year and month and cat:
            
            # Convert year and month to integers
            year = int(year)
            month = int(month)

            # Calculate the start and end dates for the selected month and year
            start_date = datetime(year, month, 1).date()
            end_date = (start_date + timedelta(days=32)).replace(day=1) - timedelta(days=1)

            # Fetch all active trainees
            if cat == 'all':all_trainees = Trainer.objects.filter(is_active=True,organization=org)
            else:all_trainees = Trainer.objects.filter(is_active=True,category=cat,organization=org)

            # Fetch trainees who have made ANY payment within the selected month
            paid_trainees_ids = Payments.objects.filter(
                paymentdate__range=(start_date, end_date),
                paymentCategry='month',  # Assuming 'month' is the category for monthly payments
                organization=org
            ).values_list('trainer_id', flat=True).distinct()

            # Exclude trainees who have made ANY payment during the selected month
            unpaid_trainees = all_trainees.exclude(id__in=paid_trainees_ids)

            # Fetch the last payment date for each unpaid trainee
            last_payment_dates = {}
            for trainee in unpaid_trainees:
                last_payment = Payments.objects.filter(trainer=trainee, paymentCategry='month', organization=org).order_by('-paymentdate').first()
                last_payment_dates[trainee] = last_payment.paymentdate if last_payment else None

            # Attach last payment date to each trainee
            for trainee in unpaid_trainees:
                trainee.last_payment_date = last_payment_dates.get(trainee, None)

            return render(request, 'pages/unpaid_trainees.html', {
                'unpaid_trainees': unpaid_trainees,
                'year': year,
                'month': month,
                'years': years,
                'cat': cat,
                'months': months,
                'selected_year': year,
                'selected_month': month,
            })

    return render(request, 'pages/unpaid_trainees.html', {
        'years': years,
        'months': months,
    })

import openpyxl  
# adjust to your actual model import

#upload trainers from excel
@login_required(login_url='/login/')
@require_organization
def upload_trainers_excel(request):
    org = request.organization
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']

        try:
            wb = openpyxl.load_workbook(excel_file)
            sheet = wb.active

            added_count = 0
            for i, row in enumerate(sheet.iter_rows(min_row=2), start=2):  # Skip header
                try:
                    first_name = row[0].value
                    last_name = row[1].value
                    birthday = row[2].value
                    gender = row[3].value
                    phone = row[4].value or 0
                    phone_parent = row[5].value or 0
                    email = row[6].value or "None@email.com"
                    address = row[7].value or org.location
                    cin = row[8].value or "لا يوجد"
                    education = row[9].value
                    belt = row[10].value
                    category = row[11].value
                    height = row[12].value or 0
                    weight = row[13].value or 0

                    if first_name and last_name and birthday and gender and education and category:
                        Trainer.objects.create(
                            organization = org,
                            first_name=first_name,
                            last_name=last_name,
                            birth_day=birthday,
                            phone=phone,
                            phone_parent=phone_parent,
                            email=email,
                            address=address,
                            CIN=cin,
                            male_female=gender,
                            belt_degree=belt,
                            Degree=education,
                            category=category,
                            started_day=datetime.today(),
                            tall=height,
                            weight=weight
                        )
                        added_count += 1
                except Exception as row_err:
                    messages.warning(request, f"خطأ في الصف {i}: {str(row_err)}")

            messages.success(request, f"تمت إضافة {added_count} مدرب(ة) بنجاح.")
            return redirect('upload_trainers_excel')  # replace with your actual URL name

        except Exception as e:
            messages.error(request, f"حدث خطأ أثناء قراءة الملف: {str(e)}")

    return render(request, "pages/upload_trainers.html")

@login_required(login_url='/login/')
@require_organization
def upload_payments_excel(request):

    org = request.organization
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']

        try:
            wb = openpyxl.load_workbook(excel_file)
            sheet = wb.active

            added = 0
            for i, row in enumerate(sheet.iter_rows(min_row=2), start=2):  # skip headers
                try:
                    trainer_id = row[0].value
                    payment_date = row[1].value or datetime.today()
                    payment_category = row[2].value or 'month'
                    payment_amount = row[3].value or 0

                    trainer = Trainer.objects.get(id=trainer_id)

                    Payments.objects.create(
                        organization=org,
                        trainer=trainer,
                        paymentdate=payment_date,
                        paymentCategry=payment_category,
                        paymentAmount=payment_amount
                    )
                    added += 1
                except Exception as e:
                    messages.warning(request, f"⚠️ خطأ في الصف {i}: {str(e)}")

            messages.success(request, f"✅ تمت إضافة {added} دفعة بنجاح.")
            return redirect('upload_payments_excel')  # change to your actual URL name

        except Exception as e:
            messages.error(request, f"فشل رفع الملف: {str(e)}")

    return render(request, "pages/upload_payments.html")


from django.utils.text import slugify

def setup_organization(request):
    
    """
    Organization setup for new users
    This view allows users to create their organization
    """
    # Check if user already has an organization
    if request.user.is_authenticated:
        try:
            staff = Staff.objects.get(user=request.user)
            # User already has organization, redirect to home
            messages.info(request, "لديك جمعية بالفعل")
            return redirect('home')
        except Staff.DoesNotExist:
            pass
    
    if request.method == 'POST':
        name = request.POST.get('name')
        slug = request.POST.get('slug')
        email = request.POST.get('email', '')
        phone = request.POST.get('phone', '')
        
        # Validate inputs
        if not name or not slug:
            messages.error(request, "الرجاء إدخال اسم الجمعية والمعرف")
            return render(request, 'pages/setup_organization.html')
        
        # Clean slug
        slug = slugify(slug)
        
        # Check if slug already exists
        if OrganizationInfo.objects.filter(slug=slug).exists():
            messages.error(request, "هذا المعرف مستخدم بالفعل، الرجاء اختيار معرف آخر")
            return render(request, 'pages/setup_organization.html', {
                'name': name,
                'email': email,
                'phone': phone
            })
        
        try:
            # Create organization
            org = OrganizationInfo.objects.create(
                name=name,
                slug=slug,
                email=email,
                phone_number=phone,
                is_active=True,
                subscription_tier='free',
                max_trainers=50
            )
            
            # Create staff record for current user
            Staff.objects.create(
                organization=org,
                user=request.user,
                role='مالك',
                is_admin=True,
                salary=0
            )
            
            messages.success(request, f"تم إنشاء جمعية {name} بنجاح!")
            
            # Redirect based on user type
            staff = Staff.objects.get(user=request.user)
            if staff.is_admin:
                return redirect('home')
            else:
                return redirect('dashboard')
                
        except Exception as e:
            messages.error(request, f"حدث خطأ أثناء إنشاء الجمعية: {str(e)}")
            return render(request, 'pages/setup_organization.html')
    
    # GET request - show form
    return render(request, 'pages/setup_organization.html')


# trainers/views.py - Add these views

from django.db import transaction
from django.utils.text import slugify
from .models import OrganizationInfo, Staff
import re
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods

@require_http_methods(["GET"])
def check_username(request):
    """
    API endpoint to check if username is available
    Usage: /api/check-username/?username=johndoe
    """
    username = request.GET.get('username', '').strip()
    
    if not username:
        return JsonResponse({
            'available': False,
            'message': 'اسم المستخدم مطلوب'
        })
    
    if len(username) < 3:
        return JsonResponse({
            'available': False,
            'message': 'اسم المستخدم يجب أن يكون 3 أحرف على الأقل'
        })
    
    # Check if username exists
    if User.objects.filter(username=username).exists():
        return JsonResponse({
            'available': False,
            'message': 'اسم المستخدم موجود بالفعل'
        })
    
    return JsonResponse({
        'available': True,
        'message': 'اسم المستخدم متاح ✓'
    })


@require_http_methods(["GET"])
def check_slug(request):
    """
    API endpoint to check if organization slug is available
    Usage: /api/check-slug/?slug=my-org
    """
    from .models import OrganizationInfo  # Import your model
    
    slug = request.GET.get('slug', '').strip()
    
    if not slug:
        return JsonResponse({
            'available': False,
            'message': 'المعرف مطلوب'
        })
    
    if len(slug) < 3:
        return JsonResponse({
            'available': False,
            'message': 'المعرف يجب أن يكون 3 أحرف على الأقل'
        })
    
    # Check format
    import re
    if not re.match(r'^[a-z0-9-]+$', slug):
        return JsonResponse({
            'available': False,
            'message': 'المعرف يجب أن يحتوي على حروف إنجليزية صغيرة وأرقام وشرطات فقط'
        })
    
    # Check if slug exists
    if OrganizationInfo.objects.filter(slug=slug).exists():
        return JsonResponse({
            'available': False,
            'message': 'المعرف مستخدم بالفعل'
        })
    
    return JsonResponse({
        'available': True,
        'message': 'المعرف متاح ✓'
    })


def signup(request):
    """
    Complete signup flow:
    1. Create user account
    2. Create organization
    3. Link user to organization via Staff
    4. Auto-login and redirect to dashboard
    """
    if request.user.is_authenticated:
        messages.info(request, "أنت مسجل دخول بالفعل")
        return redirect('home')
    
    if request.method == 'POST':
        # User Information
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password')
        password_confirm = request.POST.get('password_confirm')
        
        # Organization Information
        org_name = request.POST.get('org_name', '').strip()
        org_slug = request.POST.get('org_slug', '').strip()
        org_email = request.POST.get('org_email', '').strip()
        org_phone = request.POST.get('org_phone', '').strip()
        org_description = request.POST.get('org_description', '').strip()
        org_address = request.POST.get('org_address', '').strip()
        
        # Personal Information (for Staff record)
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        phone = request.POST.get('phone', '').strip()
        
        # Validation
        errors = []
        
        # Validate User Data
        if not username or len(username) < 3:
            errors.append("اسم المستخدم يجب أن يكون 3 أحرف على الأقل")
        
        if User.objects.filter(username=username).exists():
            errors.append("اسم المستخدم موجود بالفعل")
        
        if not email or '@' not in email:
            errors.append("البريد الإلكتروني غير صالح")
        
        if User.objects.filter(email=email).exists():
            errors.append("البريد الإلكتروني مستخدم بالفعل")
        
        if not password or len(password) < 6:
            errors.append("كلمة المرور يجب أن تكون 6 أحرف على الأقل")
        
        if password != password_confirm:
            errors.append("كلمات المرور غير متطابقة")
        
        # Validate Organization Data
        if not org_name or len(org_name) < 3:
            errors.append("اسم الجمعية يجب أن يكون 3 أحرف على الأقل")
        
        if not org_slug:
            org_slug = slugify(org_name)
        else:
            org_slug = slugify(org_slug)
        
        # Check slug format
        if not re.match(r'^[a-z0-9-]+$', org_slug):
            errors.append("المعرف يجب أن يحتوي على حروف إنجليزية صغيرة وأرقام وشرطات فقط")
        
        if OrganizationInfo.objects.filter(slug=org_slug).exists():
            errors.append("المعرف مستخدم بالفعل، الرجاء اختيار معرف آخر")
        
        # If there are errors, show them
        if errors:
            for error in errors:
                messages.error(request, error)
            return render(request, 'pages/signup.html', {
                'username': username,
                'email': email,
                'org_name': org_name,
                'org_slug': org_slug,
                'org_email': org_email,
                'org_phone': org_phone,
                'org_description': org_description,
                'org_address': org_address,
                'first_name': first_name,
                'last_name': last_name,
                'phone': phone,
            })
        
        # All validation passed - Create everything in a transaction
        try:
            with transaction.atomic():
                # 1. Create User
                user = User.objects.create_user(
                    username=username,
                    email=email,
                    password=password,
                    first_name=first_name,
                    last_name=last_name
                )
                
                # 2. Create Organization
                organization = OrganizationInfo.objects.create(
                    name=org_name,
                    slug=org_slug,
                    email=org_email or email,  # Use user email if org email not provided
                    phone_number=org_phone,
                    description=org_description,
                    is_active=True,
                    subscription_tier='free',
                    max_trainers=50  # Default for new organizations
                )
                
                # 3. Create Staff record (link user to organization)
                Staff.objects.create(
                    organization=organization,
                    user=user,
                    role='مالك',  # Owner role
                    is_admin=True,
                    salary=0,
                    email=email,
                    phone_number=phone
                )
                
                # 4. Auto-login the user
                login(request, user)
                
                # 5. Success message and redirect
                messages.success(
                    request,
                    f'مرحباً {first_name}! تم إنشاء حسابك وجمعية {org_name} بنجاح'
                )
                
                return redirect('home')  # or 'dashboard' or 'welcome'
                
        except Exception as e:
            messages.error(request, f'حدث خطأ أثناء التسجيل: {str(e)}')
            return render(request, 'pages/signup.html')
    
    # GET request - show form
    return render(request, 'pages/signup.html')


def check_username_availability(request):
    """
    AJAX endpoint to check if username is available
    """
    username = request.GET.get('username', '')
    
    if len(username) < 3:
        return JsonResponse({
            'available': False,
            'message': 'اسم المستخدم يجب أن يكون 3 أحرف على الأقل'
        })
    
    available = not User.objects.filter(username=username).exists()
    
    return JsonResponse({
        'available': available,
        'message': 'اسم المستخدم متاح' if available else 'اسم المستخدم مستخدم بالفعل'
    })


def check_slug_availability(request):
    """
    AJAX endpoint to check if organization slug is available
    """
    slug = request.GET.get('slug', '')
    slug = slugify(slug)
    
    if not slug or not re.match(r'^[a-z0-9-]+$', slug):
        return JsonResponse({
            'available': False,
            'message': 'المعرف غير صالح'
        })
    
    available = not OrganizationInfo.objects.filter(slug=slug).exists()
    
    return JsonResponse({
        'available': available,
        'message': 'المعرف متاح' if available else 'المعرف مستخدم بالفعل',
        'slug': slug
    })


def welcome(request):
    """
    Welcome page after signup with quick start guide
    """
    if not request.user.is_authenticated:
        return redirect('login')
    
    try:
        staff = Staff.objects.get(user=request.user)
        organization = staff.organization
    except Staff.DoesNotExist:
        return redirect('setup_organization')
    
    # Get registration URL for their organization
    registration_url = request.build_absolute_uri(
        reverse('addme', kwargs={'org_slug': organization.slug})
    )
    
    # Check if this is their first login
    is_first_login = request.session.get('first_login', False)
    if is_first_login:
        del request.session['first_login']
    
    context = {
        'organization': organization,
        'registration_url': registration_url,
        'is_first_login': is_first_login,
        'trainer_count': organization.trainers.filter(is_active=True).count(),
        'payment_count': organization.payments.count(),
    }
    
    return render(request, 'pages/welcome.html', context)


# Alternative: Wizard-style signup (multi-step)

def signup_step1(request):
    """Step 1: User Account"""
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        
        # Validate
        if User.objects.filter(username=username).exists():
            messages.error(request, 'اسم المستخدم موجود بالفعل')
            return render(request, 'pages/signup_step1.html')
        
        # Store in session
        request.session['signup_data'] = {
            'username': username,
            'email': email,
            'password': password,
            'first_name': first_name,
            'last_name': last_name,
        }
        
        return redirect('signup_step2')
    
    return render(request, 'pages/signup_step1.html')


def signup_step2(request):
    """Step 2: Organization Information"""
    if 'signup_data' not in request.session:
        return redirect('signup_step1')
    
    if request.method == 'POST':
        org_name = request.POST.get('org_name')
        org_slug = slugify(request.POST.get('org_slug', org_name))
        org_email = request.POST.get('org_email')
        org_phone = request.POST.get('org_phone')
        
        # Validate
        if OrganizationInfo.objects.filter(slug=org_slug).exists():
            messages.error(request, 'المعرف مستخدم بالفعل')
            return render(request, 'pages/signup_step2.html')
        
        # Get user data from session
        user_data = request.session['signup_data']
        
        # Create everything
        try:
            with transaction.atomic():
                # Create user
                user = User.objects.create_user(
                    username=user_data['username'],
                    email=user_data['email'],
                    password=user_data['password'],
                    first_name=user_data['first_name'],
                    last_name=user_data['last_name']
                )
                
                # Create organization
                organization = OrganizationInfo.objects.create(
                    name=org_name,
                    slug=org_slug,
                    email=org_email,
                    phone_number=org_phone,
                    is_active=True,
                    subscription_tier='free',
                    max_trainers=50
                )
                
                # Create staff
                Staff.objects.create(
                    organization=organization,
                    user=user,
                    role='مالك',
                    is_admin=True,
                    salary=0
                )
                
                # Clear session
                del request.session['signup_data']
                
                # Login
                login(request, user)
                request.session['first_login'] = True
                
                return redirect('welcome')
                
        except Exception as e:
            messages.error(request, f'حدث خطأ: {str(e)}')
    
    return render(request, 'pages/signup_step2.html')


# Social Auth Integration (Optional)

def signup_with_google(request):
    """
    Google OAuth signup - creates organization during callback
    Requires: pip install social-auth-app-django
    """
    # This would be configured with Django Social Auth
    # User signs in with Google, then redirected to organization setup
    pass


def complete_social_signup(request):
    """
    Complete signup after social auth
    User already created, just need to create organization
    """
    if not request.user.is_authenticated:
        return redirect('login')
    
    # Check if user already has organization
    try:
        staff = Staff.objects.get(user=request.user)
        return redirect('home')
    except Staff.DoesNotExist:
        pass
    
    if request.method == 'POST':
        org_name = request.POST.get('org_name')
        org_slug = slugify(request.POST.get('org_slug', org_name))
        
        # Validate
        if OrganizationInfo.objects.filter(slug=org_slug).exists():
            messages.error(request, 'المعرف مستخدم بالفعل')
            return render(request, 'pages/complete_social_signup.html')
        
        # Create organization and staff
        try:
            with transaction.atomic():
                organization = OrganizationInfo.objects.create(
                    name=org_name,
                    slug=org_slug,
                    email=request.user.email,
                    is_active=True,
                    subscription_tier='free',
                    max_trainers=50
                )
                
                Staff.objects.create(
                    organization=organization,
                    user=request.user,
                    role='مالك',
                    is_admin=True,
                    salary=0,
                    email=request.user.email
                )
                
                messages.success(request, f'تم إنشاء {org_name} بنجاح')
                return redirect('welcome')
                
        except Exception as e:
            messages.error(request, f'حدث خطأ: {str(e)}')
    
    return render(request, 'pages/complete_social_signup.html')